#!/usr/bin/env python3
"""
Script to upload all .docx files from the data directory to Atlassian Confluence
while preserving the folder structure.

Uses the Confluence REST API v2 to create pages from DOCX files:
https://developer.atlassian.com/cloud/confluence/rest/v2/intro/
"""

import os
import sys
import base64
import mimetypes
import requests
import json
import docx
from bs4 import BeautifulSoup
from pathlib import Path
from urllib.parse import quote
import re

# Load configuration from environment variables
import os
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

# Get configuration from environment variables
CONFLUENCE_BASE_URL = os.getenv('CONFLUENCE_BASE_URL')
API_TOKEN = os.getenv('API_TOKEN')
USERNAME = os.getenv('USERNAME')
SPACE_KEY = os.getenv('SPACE_KEY')
ROOT_PAGE_ID = os.getenv('ROOT_PAGE_ID')
SPACE_ID = None  # Will be set after retrieving numeric ID

# Get organization-wide group from config for internal permissions
ORG_GROUP = os.getenv('ORG_GROUP', 'confluence-users')  # Default group for [INT] permissions

# Permission configuration - these will be detected from filenames
# [INT] = Internal (organization-wide)
# [PUB] = Public (anyone with link)
# [RES] = Restricted (owner/explicitly shared)
INTERNAL_SUFFIX = '[INT]'
PUBLIC_SUFFIX = '[PUB]'
RESTRICTED_SUFFIX = '[RES]'

# Ensure trailing slash for base URL
if CONFLUENCE_BASE_URL and not CONFLUENCE_BASE_URL.endswith('/'):
    CONFLUENCE_BASE_URL += '/'

def get_auth_header():
    """Create the authentication header for Confluence API calls."""
    auth_str = f"{USERNAME}:{API_TOKEN}"
    auth_bytes = auth_str.encode('ascii')
    base64_bytes = base64.b64encode(auth_bytes)
    base64_auth = base64_bytes.decode('ascii')
    return {"Authorization": f"Basic {base64_auth}"}

def get_space_id(space_key):
    """Get the numeric space ID from the space key.
    
    Args:
        space_key (str): The space key to look up
        
    Returns:
        int: The numeric space ID, or None if not found
    """
    # API endpoint for getting space info
    url = f"{CONFLUENCE_BASE_URL}wiki/api/v2/spaces"
    
    params = {"keys": space_key}
    headers = get_auth_header()
    
    try:
        response = requests.get(url, headers=headers, params=params)
        response.raise_for_status()
        results = response.json().get("results", [])
        if results:
            space_id = results[0].get("id")
            print(f"Found space ID {space_id} for space key {space_key}")
            return space_id
        else:
            print(f"Error: Space with key '{space_key}' not found")
            return None
    except requests.exceptions.RequestException as e:
        print(f"Error retrieving space ID for '{space_key}': {e}")
        if hasattr(e, 'response') and e.response is not None:
            print(f"Response: {e.response.text}")
        return None

def create_page(title, parent_id=None, space_id=None, is_folder=True):
    """
    Create a new page in Confluence.
    
    Args:
        title (str): Title of the page
        parent_id (str): ID of the parent page, or None for root page
        space_id (int): Numeric space ID (not space key)
        is_folder (bool): If True, create a simple folder page; otherwise create a normal page
    
    Returns:
        str: ID of the created page, or None if failed
    """
    global SPACE_ID
    
    if space_id is None:
        space_id = SPACE_ID
    
    # Check if a page with this title already exists
    existing_page_id = find_page_by_title(title, space_id, parent_id)
    if existing_page_id:
        print(f"Found existing page: {title} (ID: {existing_page_id})")
        return existing_page_id
        
    # Use the v2 API endpoint
    url = f"{CONFLUENCE_BASE_URL}wiki/api/v2/pages"
    
    # Content for folder pages vs. regular pages
    content = f"<p>Folder: {title}</p>" if is_folder else f"<p>Page: {title}</p>"
    
    # Create page content
    data = {
        "spaceId": space_id,
        "status": "current",
        "title": title,
        "body": {
            "representation": "storage",
            "value": content
        }
    }
    
    # Add parent reference if provided
    if parent_id:
        data["parentId"] = parent_id
    
    headers = {
        **get_auth_header(),
        "Content-Type": "application/json"
    }
    
    try:
        response = requests.post(url, headers=headers, data=json.dumps(data))
        response.raise_for_status()
        page_data = response.json()
        print(f"Created new page: {title} (ID: {page_data['id']})")
        return page_data["id"]
    except requests.exceptions.RequestException as e:
        print(f"Error creating page '{title}': {e}")
        if hasattr(e, 'response') and e.response is not None:
            print(f"Response: {e.response.text}")
        return None

def find_page_by_title(title, space_id=None, parent_id=None):
    """
    Find a page by title in a specific space and optionally under a specific parent.
    
    Args:
        title (str): Title of the page to find
        space_id (int): Numeric space ID (not space key)
        parent_id (str): Optional parent page ID
    
    Returns:
        str: ID of the page if found, None otherwise
    """
    global SPACE_ID
    
    if space_id is None:
        space_id = SPACE_ID
    
    # Use the v2 API endpoint
    url = f"{CONFLUENCE_BASE_URL}wiki/api/v2/pages"
    
    # Build query parameters
    params = {
        "title": title,
        "spaceId": space_id,
        "status": "current",
        "limit": 10
    }
    
    if parent_id:
        params["parentId"] = parent_id
    
    headers = get_auth_header()
    
    try:
        response = requests.get(url, headers=headers, params=params)
        response.raise_for_status()
        results = response.json()["results"]
        
        if results:
            page_id = results[0]["id"]
            return page_id
        return None
    except requests.exceptions.RequestException as e:
        print(f"Error finding page with title '{title}': {e}")
        if hasattr(e, 'response') and e.response is not None:
            print(f"Response: {e.response.text}")
        return None

def get_or_create_page(title, parent_id=None):
    """
    Get a page by title or create it if it doesn't exist.
    
    Args:
        title (str): Title of the page
        parent_id (str): Optional parent page ID
    
    Returns:
        str: ID of the page (existing or newly created)
    """
    page_id = find_page_by_title(title, parent_id=parent_id)
    if page_id:
        print(f"Found existing page: {title} (ID: {page_id})")
        return page_id
    
    # Page doesn't exist, create it
    page_id = create_page(title, parent_id)
    if page_id:
        print(f"Created new page: {title} (ID: {page_id})")
    return page_id

def convert_docx_to_html(file_path):
    """
    Convert a DOCX file to HTML for Confluence.
    
    Args:
        file_path (str): Path to the DOCX file
        
    Returns:
        str: HTML content extracted from the DOCX file
    """
    try:
        doc = docx.Document(file_path)
        full_html = []
        
        # Process each paragraph
        for para in doc.paragraphs:
            # Skip empty paragraphs
            if not para.text.strip():
                continue
                
            # Determine if this is a heading
            if para.style.name.startswith('Heading'):
                heading_level = int(para.style.name.split(' ')[1])
                full_html.append(f"<h{heading_level}>{para.text}</h{heading_level}>")
            else:
                # Process paragraph text with styling
                para_html = "<p>"
                for run in para.runs:
                    text = run.text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                    if run.bold:
                        text = f"<strong>{text}</strong>"
                    if run.italic:
                        text = f"<em>{text}</em>"
                    if run.underline:
                        text = f"<u>{text}</u>"
                    para_html += text
                para_html += "</p>"
                full_html.append(para_html)
        
        # Process tables
        for table in doc.tables:
            table_html = "<table><tbody>"
            for row in table.rows:
                table_html += "<tr>"
                for cell in row.cells:
                    table_html += f"<td>{cell.text}</td>"
                table_html += "</tr>"
            table_html += "</tbody></table>"
            full_html.append(table_html)
            
        # Join all HTML elements
        return "\n".join(full_html)
    except Exception as e:
        print(f"Error converting DOCX to HTML: {e}")
        return f"<p>Error converting DOCX: {e}</p>"

def upload_docx_as_page(file_path, parent_id=None, space_id=None):
    """
    Upload a DOCX file as a Confluence page.
    
    Args:
        file_path (str): Path to the DOCX file
        parent_id (str): ID of the parent page, or None for root page
        space_id (int): Numeric space ID (not space key)
        
    Returns:
        str: ID of the created page, or None if failed
    """
    global SPACE_ID
    
    if space_id is None:
        space_id = SPACE_ID
        
    # Get the file name and detect permission level
    file_name = os.path.basename(file_path)
    
    # Detect permission level from filename suffix
    permission_level, group_name = get_permission_level_from_filename(file_name)
    
    # Get title, keeping permission suffix if present
    page_title = os.path.splitext(file_name)[0]
    
    # Check if a page with this title already exists
    existing_page_id = find_page_by_title(page_title, space_id, parent_id)
    if existing_page_id:
        # A page with this title already exists
        print(f"Page with title '{page_title}' already exists with ID {existing_page_id}")
        print(f"Updating existing page content instead of creating a new one")
        
        # Update the existing page content with the detected permissions
        return update_page_content(existing_page_id, page_title, convert_docx_to_html(file_path), 
                              permission_level, group_name)
    
    # Convert DOCX to HTML
    html_content = convert_docx_to_html(file_path)
    
    # Use Confluence REST API v2 for page creation
    url = f"{CONFLUENCE_BASE_URL}wiki/api/v2/pages"
    
    # Prepare request body
    data = {
        "spaceId": space_id,
        "status": "current",
        "title": page_title,
        "body": {
            "representation": "storage",
            "value": html_content
        }
    }
    
    # Add parent if provided
    if parent_id:
        data["parentId"] = parent_id
        
    headers = {
        **get_auth_header(),
        "Content-Type": "application/json"
    }
    
    try:
        response = requests.post(url, headers=headers, data=json.dumps(data))
        response.raise_for_status()
        page_data = response.json()
        print(f"Successfully created page '{page_title}' with ID {page_data['id']}")
        
        # Upload the original document as an attachment
        if upload_attachment_to_page(file_path, page_data['id']):
            print(f"Uploaded original document as attachment to page: {page_title}")
        else:
            print(f"Failed to upload original document as attachment to page: {page_title}")
        
        # Apply permissions based on detected level from filename
        if permission_level:
            if apply_permissions_by_level(page_data['id'], page_title, permission_level, group_name):
                print(f"Applied {permission_level} permissions to page: {page_title}")
            else:
                print(f"Failed to apply {permission_level} permissions to page: {page_title}")
        
        return page_data["id"]
    except requests.exceptions.RequestException as e:
        print(f"Error creating page '{page_title}': {e}")
        if hasattr(e, 'response') and e.response is not None:
            print(f"Response: {e.response.text}")
        return None
        
def get_permission_level_from_filename(file_name):
    """
    Determine the permission level based on file name suffix.
    
    Args:
        file_name (str): Name of the file to check
    
    Returns:
        tuple: (permission_type, group_name) where permission_type is one of:
               'public', 'internal', 'restricted', or None if no suffix found
    """
    # Default to None (no restrictions)
    permission_level = None
    group_name = None
    
    # Check for each suffix
    if file_name.endswith(INTERNAL_SUFFIX + '.docx'):
        permission_level = 'internal'
        group_name = ORG_GROUP
    elif file_name.endswith(PUBLIC_SUFFIX + '.docx'):
        permission_level = 'public'
        # No group needed for public
        group_name = None
    elif file_name.endswith(RESTRICTED_SUFFIX + '.docx'):
        permission_level = 'restricted'
        # Will use the default space permissions
        # Could be enhanced later to use specific groups
        group_name = None
    
    return (permission_level, group_name)

def check_group_exists(group_name):
    """
    Check if a group exists in Confluence
    
    Args:
        group_name (str): Name of the group to check
        
    Returns:
        bool: True if the group exists, False otherwise
    """
    if not group_name:
        return False
        
    # First try the v2 API
    v2_url = f"{CONFLUENCE_BASE_URL}wiki/api/v2/groups/{quote(group_name)}"
    headers = get_auth_header()
    
    try:
        print(f"Checking if group exists: '{group_name}' using URL: {v2_url}")
        v2_response = requests.get(v2_url, headers=headers)
        
        if v2_response.status_code == 200:
            print(f"Group '{group_name}' found using v2 API")
            return True
        else:
            print(f"Group '{group_name}' not found using v2 API (status: {v2_response.status_code})")
            print(f"Response: {v2_response.text}")
            
            # Fallback to v1 API
            v1_url = f"{CONFLUENCE_BASE_URL}wiki/rest/api/group/{quote(group_name)}"
            print(f"Trying v1 API URL: {v1_url}")
            v1_response = requests.get(v1_url, headers=headers)
            
            if v1_response.status_code == 200:
                print(f"Group '{group_name}' found using v1 API")
                return True
            else:
                print(f"Group '{group_name}' not found using v1 API (status: {v1_response.status_code})")
                print(f"Response: {v1_response.text}")
                
                # Try user search API to see if the group might be visible there
                search_url = f"{CONFLUENCE_BASE_URL}wiki/rest/api/search?cql=type=group AND title~\"{group_name}\""
                print(f"Trying search API: {search_url}")
                search_response = requests.get(search_url, headers=headers)
                
                if search_response.status_code == 200:
                    results = search_response.json().get("results", [])
                    if results:
                        print(f"Found similar groups via search: {[r.get('title') for r in results]}")
                    else:
                        print(f"No similar groups found via search")
                
                return False
    except requests.exceptions.RequestException as e:
        print(f"Error checking if group exists: {e}")
        if hasattr(e, 'response') and e.response is not None:
            print(f"Response: {e.response.text}")
        return False

def apply_permissions_by_level(page_id, title, permission_level, group_name=None):
    """
    Apply appropriate permissions based on the determined level.
    
    Args:
        page_id (str): ID of the page to set permissions on
        title (str): Title of the page (for logging)
        permission_level (str): One of 'public', 'internal', 'restricted', or None
        group_name (str, optional): Not used since internal is default
    
    Returns:
        bool: True if successful, False otherwise
    """
    if not permission_level or permission_level == 'internal' or permission_level == 'public':
        # Both internal and public documents are restricted to organization members only
        # This is the default in this Confluence instance - no need to change anything
        print(f"Using default organization-only permissions for '{title}'")
        return True
            
    elif permission_level == 'restricted':
        # For restricted documents, we need to set owner-only access
        try:
            print(f"Setting '{title}' as restricted (accessible only to owner)")
            
            # First, we need to get the account ID of the current user
            # This is necessary because Atlassian Cloud APIs require accountId instead of username
            current_user_url = f"{CONFLUENCE_BASE_URL}wiki/rest/api/user/current"
            headers = get_auth_header()
            
            user_response = requests.get(current_user_url, headers=headers)
            if user_response.status_code != 200:
                print(f"Failed to get current user info: {user_response.status_code}")
                print(f"Response: {user_response.text}")
                return False
                
            user_data = user_response.json()
            account_id = user_data.get('accountId')
            
            if not account_id:
                print("Failed to get account ID for current user")
                return False
                
            print(f"Retrieved account ID: {account_id} for current user")
            
            # PUT request to replace all existing restrictions
            url = f"{CONFLUENCE_BASE_URL}wiki/rest/api/content/{page_id}/restriction"
            headers["Content-Type"] = "application/json"
            
            # Define the payload for both read and update restrictions
            # Using accountId instead of username
            payload = [
                {
                    "operation": "read",
                    "restrictions": {
                        "user": [
                            {
                                "type": "known",
                                "accountId": account_id
                            }
                        ]
                    }
                },
                {
                    "operation": "update",
                    "restrictions": {
                        "user": [
                            {
                                "type": "known",
                                "accountId": account_id
                            }
                        ]
                    }
                }
            ]
            
            response = requests.put(url, headers=headers, data=json.dumps(payload))
            
            if response.status_code >= 200 and response.status_code < 300:
                print(f"Successfully set restricted permissions for '{title}'")
                return True
            else:
                print(f"Failed to set restricted permissions for '{title}': {response.status_code}")
                print(f"Response: {response.text}")
                return False
        except Exception as e:
            print(f"Error setting restricted permissions for '{title}': {e}")
            return False
            
    return False  # Should never get here
        
    return False

def set_page_restrictions(page_id, restriction_type, group_name):
    """
    Set restrictions on a Confluence page for a specific group.
    
    Args:
        page_id (str): ID of the page to restrict
        restriction_type (str): Type of restriction ('read' or 'update')
        group_name (str): Name of the group to grant access to
        
    Returns:
        bool: True if successful, False otherwise
    """
    if not (page_id and restriction_type and group_name):
        print(f"Skipping restrictions for page {page_id} (missing parameters)")
        return False
    
    try:
        headers = get_auth_header()
        headers["Content-Type"] = "application/json"
        
        # Try the v2 API first (most reliable in newer Confluence Cloud)
        v2_url = f"{CONFLUENCE_BASE_URL}wiki/api/v2/pages/{page_id}/permissions"  # Updated URL pattern for v2 API
        v2_payload = {
            "operationType": "addPermission",
            "subject": {
                "type": "group",
                "identifier": group_name
            },
            "operation": {
                "key": restriction_type,
                "targetType": "page"  # Updated to specify 'page' instead of generic 'content'
            }
        }
        
        print(f"Attempting v2 API permission call to URL: {v2_url}")
        print(f"Payload: {json.dumps(v2_payload)}")
        v2_response = requests.post(v2_url, headers=headers, json=v2_payload)
        print(f"V2 API response status: {v2_response.status_code}")
        print(f"V2 API response: {v2_response.text}")
        
        if v2_response.status_code >= 200 and v2_response.status_code < 300:
            print(f"Successfully restricted {restriction_type} access on page {page_id} to group '{group_name}' using v2 API")
            return True
        else:
            print(f"V2 API failed with status {v2_response.status_code}, trying v1 API...")
            
            # Try the v1 API next
            v1_url = f"{CONFLUENCE_BASE_URL}wiki/rest/api/content/{page_id}/restriction/{restriction_type}"
            
            # Create the restriction payload for v1 API
            v1_payload = {
                "user": [],
                "group": [
                    group_name
                ]
            }
            
            # Create the new restriction
            v1_response = requests.post(v1_url, headers=headers, json=v1_payload)
            
            # Check status code directly - some Confluence instances return non-standard codes
            if v1_response.status_code >= 200 and v1_response.status_code < 300:
                print(f"Successfully restricted {restriction_type} access on page {page_id} to group '{group_name}' using v1 API")
                return True
            else:
                # If both v2 and v1 failed, try the experimental API as last resort
                print(f"V1 API also failed, trying experimental API...")
                
                # Experimental endpoint for setting permissions
                exp_url = f"{CONFLUENCE_BASE_URL}wiki/rest/experimental/content/{page_id}/restriction"
                
                exp_payload = {
                    "restrictions": {
                        restriction_type: {
                            "group": [group_name]
                        }
                    }
                }
                
                exp_response = requests.put(exp_url, headers=headers, json=exp_payload)
                if exp_response.status_code >= 200 and exp_response.status_code < 300:
                    print(f"Successfully restricted {restriction_type} access on page {page_id} using experimental API")
                    return True
                else:
                    print(f"All permission APIs failed. Last error: {exp_response.status_code} {exp_response.reason}")
                    if exp_response.text:
                        print(f"Response: {exp_response.text}")
                    return False
        
    except requests.exceptions.RequestException as e:
        print(f"Error setting page restrictions: {e}")
        if hasattr(e, 'response') and e.response is not None:
            print(f"Response: {e.response.text}")
        return False

def remove_all_restrictions(page_id):
    """
    Remove all restrictions from a page to make it accessible to everyone with space access.
    
    Args:
        page_id (str): ID of the page to remove restrictions from
        
    Returns:
        bool: True if successful, False otherwise
    """
    if not page_id:
        print("Cannot remove restrictions (missing page ID)")
        return False
        
    try:
        headers = get_auth_header()
        headers["Content-Type"] = "application/json"
        
        # First try the standard API endpoint for restriction deletion
        for restriction_type in ["read", "update"]:
            url = f"{CONFLUENCE_BASE_URL}wiki/rest/api/content/{page_id}/restriction/{restriction_type}"
            
            # Get current restrictions to see if any exist
            get_response = requests.get(url, headers=headers)
            if get_response.status_code >= 200 and get_response.status_code < 300:
                restrictions_data = get_response.json()
                if "results" in restrictions_data and len(restrictions_data["results"]) > 0:
                    # Restrictions exist, delete them
                    delete_response = requests.delete(url, headers=headers)
                    if delete_response.status_code < 200 or delete_response.status_code >= 300:
                        print(f"Failed to remove {restriction_type} restrictions: {delete_response.status_code} {delete_response.reason}")
                        # Try the experimental API as fallback
                        exp_url = f"{CONFLUENCE_BASE_URL}wiki/rest/experimental/content/{page_id}/restriction"
                        exp_payload = {"restrictions": {restriction_type: {"user": [], "group": []}}}
                        exp_response = requests.put(exp_url, headers=headers, json=exp_payload)
                        if exp_response.status_code < 200 or exp_response.status_code >= 300:
                            print(f"Failed to remove {restriction_type} restrictions with experimental API: {exp_response.status_code}")
                            return False
        
        print(f"Successfully removed all restrictions from page {page_id}")
        return True
    except requests.exceptions.RequestException as e:
        print(f"Error removing page restrictions: {e}")
        if hasattr(e, 'response') and e.response is not None:
            print(f"Response: {e.response.text}")
        return False

def enable_anonymous_access(page_id):
    """
    Enable anonymous access to a page if the Confluence instance supports it.
    This makes the page truly public (anyone with link can access).
    
    Args:
        page_id (str): ID of the page to enable anonymous access for
        
    Returns:
        bool: True if successful, False if failed, None if API not supported
    """
    if not page_id:
        print("Cannot enable anonymous access (missing page ID)")
        return False
        
    try:
        headers = get_auth_header()
        headers["Content-Type"] = "application/json"
        
        # Try to use the space permissions API to check if anonymous access is possible
        # First we need to get the page details to find the space key/id
        page_url = f"{CONFLUENCE_BASE_URL}wiki/rest/api/content/{page_id}?expand=space"
        page_response = requests.get(page_url, headers=headers)
        
        if page_response.status_code >= 200 and page_response.status_code < 300:
            page_data = page_response.json()
            space_key = page_data.get("space", {}).get("key")
            
            if space_key:
                # Try to enable anonymous access at the page-level instead
                # Use the REST API v2 to make the content publicly accessible
                content_perm_url = f"{CONFLUENCE_BASE_URL}wiki/api/v2/content/{page_id}/permissions"
                content_perm_payload = {
                    "operationType": "addPermission",
                    "subject": {
                        "type": "anonymous"
                    },
                    "operation": {
                        "key": "read",
                        "targetType": "content"
                    }
                }
                
                content_perm_response = requests.post(content_perm_url, headers=headers, json=content_perm_payload)
                
                if content_perm_response.status_code >= 200 and content_perm_response.status_code < 300:
                    print(f"Successfully enabled anonymous access for page {page_id}")
                    return True
                else:
                    # If the v2 API fails, fall back to the space property approach
                    print(f"V2 permission API failed, falling back to space property method")
                    
                    # Try the space property API
                    anon_url = f"{CONFLUENCE_BASE_URL}wiki/rest/api/space/{space_key}/property/anonymous-access"
                    
                    # First check if property exists
                    check_response = requests.get(anon_url, headers=headers)
                    
                    if check_response.status_code == 200:
                        # Property exists, need to include version in update
                        property_data = check_response.json()
                        version = property_data.get("version", {}).get("number", 0)
                        
                        # Update existing property with version
                        anon_payload = {
                            "value": "true",
                            "version": {"number": version + 1}  # Increment version
                        }
                        anon_response = requests.put(anon_url, headers=headers, json=anon_payload)
                    elif check_response.status_code == 404:
                        # Create new property
                        anon_payload = {"value": "true", "key": "anonymous-access"}
                        anon_response = requests.post(anon_url, headers=headers, json=anon_payload)
                    else:
                        print(f"Unexpected status checking anonymous property: {check_response.status_code}")
                        return None
                        
                    if anon_response.status_code >= 200 and anon_response.status_code < 300:
                        print(f"Successfully enabled anonymous access for page in space {space_key}")
                        return True
                    else:
                        print(f"Failed to enable anonymous access: {anon_response.status_code}")
                        if anon_response.text:
                            print(f"Response: {anon_response.text}")
                        return False
            else:
                print(f"Could not determine space key for page {page_id}")
                return False
        else:
            print(f"Failed to get page details: {page_response.status_code}")
            return False
    except requests.exceptions.RequestException as e:
        print(f"Error enabling anonymous access: {e}")
        # If we get here, the API endpoint likely doesn't exist
        return None

def set_restricted_permissions(page_id):
    """
    Set restricted permissions on a page (owner only + explicit shares).
    
    Args:
        page_id (str): ID of the page to restrict
        
    Returns:
        bool: True if successful, False otherwise
    """
    if not page_id:
        print("Cannot set restricted permissions (missing page ID)")
        return False
        
    try:
        headers = get_auth_header()
        headers["Content-Type"] = "application/json"
        
        # Get the current user details to set owner-only permissions
        # Use the global USERNAME from environment variables instead of trying to fetch current user
        if USERNAME:
            # Set both read and update restrictions to the current user (owner)
            for restriction_type in ["read", "update"]:
                url = f"{CONFLUENCE_BASE_URL}wiki/rest/api/content/{page_id}/restriction/{restriction_type}"
                
                # Create payload with just the owner as allowed user
                payload = {
                    "user": [USERNAME],  # Use the configured username
                    "group": []
                }
                
                response = requests.post(url, headers=headers, json=payload)
                
                if response.status_code < 200 or response.status_code >= 300:
                    print(f"Failed to set {restriction_type} restriction to owner-only: {response.status_code}")
                    
                    # Try experimental API as fallback
                    exp_url = f"{CONFLUENCE_BASE_URL}wiki/rest/experimental/content/{page_id}/restriction"
                    exp_payload = {"restrictions": {restriction_type: {"user": [USERNAME], "group": []}}}
                    
                    exp_response = requests.put(exp_url, headers=headers, json=exp_payload)
                    if exp_response.status_code < 200 or exp_response.status_code >= 300:
                        print(f"Failed to set {restriction_type} restriction with experimental API: {exp_response.status_code}")
                        if exp_response.text:
                            print(f"Response: {exp_response.text}")
                        return False
            
            print(f"Successfully set restricted (owner-only) permissions for page {page_id}")
            return True
        else:
            print(f"Error: USERNAME not set in environment variables. Cannot set restricted permissions.")
            return False
    except requests.exceptions.RequestException as e:
        print(f"Error setting restricted permissions: {e}")
        if hasattr(e, 'response') and e.response is not None:
            print(f"Response: {e.response.text}")
        return False

def upload_attachment_to_page(file_path, page_id):
    """
    Upload a file as an attachment to a Confluence page.
    
    Args:
        file_path (str): Path to the file to upload
        page_id (str): ID of the page to attach the file to
    
    Returns:
        bool: True if successful, False otherwise
    """
    # For API v2, attachments must be created using a different endpoint
    # Using REST API v1 for attachments since v2 seems to have issues with POST
    url = f"{CONFLUENCE_BASE_URL}wiki/rest/api/content/{page_id}/child/attachment"
    
    # Prepare the file to upload
    file_name = os.path.basename(file_path)
    
    # Open the file in binary mode
    files = {'file': (file_name, open(file_path, 'rb'))}
    
    # Add the authentication header
    headers = get_auth_header()
    headers['X-Atlassian-Token'] = 'no-check' # Required for file uploads
    
    try:
        response = requests.post(url, headers=headers, files=files)
        response.raise_for_status()
        print(f"Successfully uploaded attachment '{file_name}' to page {page_id}")
        return True
    except requests.exceptions.RequestException as e:
        print(f"Error uploading attachment '{file_name}': {e}")
        if hasattr(e, 'response') and e.response is not None:
            print(f"Response: {e.response.text}")
        return False
    finally:
        # Close the file
        files['file'][1].close()

def get_page_info(page_id):
    """
    Get information about a page including its current version.
    
    Args:
        page_id (str): ID of the page to get information for
        
    Returns:
        dict: Page information including version, or None if failed
    """
    url = f"{CONFLUENCE_BASE_URL}wiki/api/v2/pages/{page_id}"
    headers = get_auth_header()
    
    try:
        response = requests.get(url, headers=headers)
        response.raise_for_status()
        return response.json()
    except requests.exceptions.RequestException as e:
        print(f"Error getting page info for ID {page_id}: {e}")
        if hasattr(e, 'response') and e.response is not None:
            print(f"Response: {e.response.text}")
        return None

def update_page_content(page_id, title, html_content, permission_level=None, group_name=None):
    """
    Update an existing Confluence page with new content.
    
    Args:
        page_id (str): ID of the page to update
        title (str): Title of the page
        html_content (str): HTML content to update the page with
        permission_level (str, optional): Permission level to apply
        group_name (str, optional): Group name for internal permission level
    
    Returns:
        str: Page ID if successful, None otherwise
    """
    # First get the current page info to get version number
    page_info = get_page_info(page_id)
    if not page_info:
        print(f"Cannot update page {page_id}: Unable to retrieve current page information")
        return None
        
    # Get the current version number
    version = page_info.get("version", {}).get("number")
    if not version:
        print(f"Cannot update page {page_id}: Unable to determine current version number")
        return None
        
    # Use the v2 API endpoint for page updates
    url = f"{CONFLUENCE_BASE_URL}wiki/api/v2/pages/{page_id}"
    
    # Prepare request body with version information
    data = {
        "id": page_id,
        "status": "current",
        "title": title,
        "body": {
            "representation": "storage",
            "value": html_content
        },
        "version": {
            "number": version + 1
        }
    }
    
    # Add the authentication header
    headers = get_auth_header()
    headers["Content-Type"] = "application/json"
    
    try:
        response = requests.put(url, headers=headers, json=data)
        response.raise_for_status()
        print(f"Successfully updated page content for '{title}' with ID {page_id}")
        
        # Apply permissions based on detected level from filename
        if permission_level:
            if apply_permissions_by_level(page_id, title, permission_level, group_name):
                print(f"Applied {permission_level} permissions to updated page: {title}")
            else:
                print(f"Failed to apply {permission_level} permissions to updated page: {title}")
                
        return page_id
    except requests.exceptions.RequestException as e:
        print(f"Error updating page '{title}': {e}")
        if hasattr(e, 'response') and e.response is not None:
            print(f"Response: {e.response.text}")
        return None

def update_folder_page_with_links(parent_folder_id, child_pages, folder_children):
    """
    Update a folder page to include links to its child pages, separated into folders and regular pages.
    
    Args:
        folder_id (str): ID of the folder page to update
        child_pages (list): List of tuples (page_title, page_id) for child pages
        folder_children (dict): Dictionary mapping folder paths to their child pages
        
    Returns:
        bool: True if successful, False otherwise
    """
    if not child_pages:
        return True
    
    # Get folder info to get title and current version
    folder_info = get_page_info(parent_folder_id)
    if not folder_info:
        return False
    
    folder_title = folder_info.get("title", "Folder")
    
    # Split child pages into folders and regular pages
    folders = []
    regular_pages = []
    
    for page_title, page_id in child_pages:
        # Check if this child is a folder by examining paths
        is_folder = False
        
        # First check if this child has its own children list
        path_to_check = None
        
        # Check both relative paths - with and without parent path prefix
        possible_paths = []
        for path in folder_children.keys():
            path_parts = path.split('/')
            if page_title in path_parts:
                if path_parts[-1] == page_title:  # It's the last part of the path
                    is_folder = True
                    break
        
        if is_folder:
            folders.append((page_title, page_id))
        else:
            regular_pages.append((page_title, page_id))
    
    # Create HTML content with links to child pages
    html_content = f"<h1>Folder: {folder_title}</h1>\n"
    
    # Add folders section if there are any folders
    if folders:
        html_content += "<h2>This folder contains the following folders:</h2>\n"
        html_content += "<ul>\n"
        for subfolder_title, subfolder_id in folders:
            html_content += f'<li><ac:link><ri:page ri:content-title="{subfolder_title}" /></ac:link></li>\n'
        html_content += "</ul>\n"
    
    # Add pages section if there are any regular pages
    if regular_pages:
        html_content += "<h2>This folder contains the following pages:</h2>\n"
        html_content += "<ul>\n"
        for page_title, page_id in regular_pages:
            html_content += f'<li><ac:link><ri:page ri:content-title="{page_title}" /></ac:link></li>\n'
        html_content += "</ul>\n"
    
    # Update the folder page with the new content
    result = update_page_content(parent_folder_id, folder_title, html_content)
    return result is not None

def upload_docx_files_to_confluence(data_dir):
    """
    Upload all .docx files from the data directory to Confluence, preserving folder structure.
    Each .docx file is converted to a Confluence page with its contents.
    
    Args:
        data_dir (str): Path to the data directory
    """
    # Start with the root page ID if provided
    parent_id_map = {"": ROOT_PAGE_ID}
    
    # Dictionary to keep track of child pages for each folder
    # Key: folder_path, Value: list of (page_title, page_id) tuples
    folder_children = {}
    
    # First pass: Create all folder pages (even if they don't have .docx files directly)
    print("Step 1: Creating folder structure...")
    for root, dirs, _ in os.walk(data_dir):
        # Get the relative path from the data directory
        rel_path = os.path.relpath(root, data_dir)
        if rel_path == '.':
            rel_path = ''
        
        # Initialize folder's child list
        if rel_path not in folder_children:
            folder_children[rel_path] = []
        
        # Split the path into components
        path_components = rel_path.split(os.sep) if rel_path else []
        
        # Build up the path and create folder pages as needed
        current_path = ""
        current_parent_id = parent_id_map[""]
        
        for component in path_components:
            if not component:
                continue
            
            # Update the current path
            if current_path:
                current_path = os.path.join(current_path, component)
            else:
                current_path = component
            
            # Check if we already have a page ID for this path
            if current_path not in parent_id_map:
                # Create the folder page and store its ID
                page_id = create_page(component, current_parent_id)
                if not page_id:
                    print(f"Failed to create page for directory: {component}")
                    break
                parent_id_map[current_path] = page_id
                
                # Add this folder as a child of its parent
                parent_path = os.path.dirname(current_path)
                if parent_path in folder_children:
                    folder_children[parent_path].append((component, page_id))
            
            # Update the current parent ID
            current_parent_id = parent_id_map[current_path]
    
    # Second pass: Upload all .docx files as pages
    print("Step 2: Creating document pages...")
    for root, _, files in os.walk(data_dir):
        # Filter for .docx files
        docx_files = [f for f in files if f.lower().endswith('.docx')]
        if not docx_files:
            continue
        
        # Get the relative path and current parent ID
        rel_path = os.path.relpath(root, data_dir)
        if rel_path == '.':
            rel_path = ''
        
        current_parent_id = parent_id_map.get(rel_path, parent_id_map[""])
        
        # Upload all .docx files in this directory as pages
        for file_name in docx_files:
            file_path = os.path.join(root, file_name)
            # Get page title from file name
            page_title = os.path.splitext(file_name)[0]
            
            # Process document file
            page_id = upload_docx_as_page(file_path, current_parent_id)
            
            if not page_id:
                print(f"Failed to upload document as page: {file_path}")
            else:
                # Add this page as a child of its parent folder
                folder_children[rel_path].append((page_title, page_id))
    
    # Third pass: Update all folder pages with links to their children
    print("Step 3: Updating folder pages with child links...")
    for folder_path, children in folder_children.items():
        # Skip empty path if ROOT_PAGE_ID is None
        if folder_path == "" and not ROOT_PAGE_ID:
            continue
            
        if folder_path in parent_id_map and children:
            folder_id = parent_id_map[folder_path]
            if folder_id:  # Make sure we have a valid folder ID
                update_folder_page_with_links(folder_id, children, folder_children)

def main():
    """Main function to run the script."""
    global SPACE_ID
    
    # Check required configuration
    if not CONFLUENCE_BASE_URL or not API_TOKEN or not USERNAME or not SPACE_KEY:
        print("Error: Please update the configuration variables in the script.")
        print("Required: CONFLUENCE_BASE_URL, API_TOKEN, USERNAME, SPACE_KEY")
        sys.exit(1)
    
    # Get the numeric space ID from the space key
    SPACE_ID = get_space_id(SPACE_KEY)
    if not SPACE_ID:
        print(f"Error: Could not find space ID for space key '{SPACE_KEY}'")
        sys.exit(1)
    
    # Define the data directory path
    data_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'data')
    
    # Check if the data directory exists
    if not os.path.isdir(data_dir):
        print(f"Error: Data directory not found at {data_dir}")
        sys.exit(1)
    
    print(f"Starting upload of .docx files from {data_dir} to Confluence...")
    upload_docx_files_to_confluence(data_dir)
    print("Upload process completed.")

if __name__ == "__main__":
    main()
